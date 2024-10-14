import os
import io
import csv
from datetime import datetime, timedelta
import pandas as pd
from pydub import AudioSegment
from pydub.silence import split_on_silence
import speech_recognition as sr
from docx import Document
import mercury as mr
from google.colab import files
def recognize_speech_google(audio_file):
    """
    Recognize speech from an audio file using Google Speech Recognition.
    """
    recognizer = sr.Recognizer()

    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)

    try:
        # Recognize the audio, assuming the language is Arabic
        text = recognizer.recognize_google(audio, language='ar-EG')
    except sr.UnknownValueError:
        text = "(Could not understand)"
    except sr.RequestError as e:
        text = f"(Error: {e})"

    return text
def main(audio_file, output_docx):
    """
    Process an audio file to recognize speech and save results to a DOCX file.
    """
    # Load the audio file
    audio = AudioSegment.from_file(audio_file)

    # Split audio into chunks based on silence
    chunks = split_on_silence(audio, min_silence_len=500, silence_thresh=-40)

    output_data = []
    start_time = datetime.now()

    # Process each chunk
    for chunk in chunks:
        # Export each chunk to a temporary wav file
        chunk.export("temp.wav", format="wav")

        # Recognize speech from the exported chunk
        text = recognize_speech_google("temp.wav")

        # Calculate timestamp for the chunk
        duration = timedelta(seconds=chunk.duration_seconds)
        end_time = start_time + duration
        timestamp = end_time.strftime("%H:%M:%S")

        # Append timestamp and recognized text to the output data
        output_data.append((timestamp, text))

        # Update start time for the next chunk
        start_time = end_time

    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Speech Recognition Output', 0)

    # Create a table with two columns (Timestamp and Text)
    table = doc.add_table(rows=1, cols=2)

    # Set the headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Timestamp'
    hdr_cells[1].text = 'Text'

    # Add the data to the table
    for timestamp, text in output_data:
        row_cells = table.add_row().cells
        row_cells[0].text = timestamp
        row_cells[1].text = text
    # Save the document
    doc.save(output_docx)

if __name__ == "__main__":
    # Upload the audio file using Google Colab's file upload widget
    print("Please upload your audio file:")
    uploaded = files.upload()

    # Get the file name from the uploaded files
    audio_file = list(uploaded.keys())[0]

    # Specify the output DOCX file name
    output_docx = "output.docx"

    # Call the main function with the uploaded audio file
    main(audio_file, output_docx)